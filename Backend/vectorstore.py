import os
import logging
import csv
import time
from typing import List
from pathlib import Path

from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from config import config

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Simplified document loader without external dependencies"""
    
    def load_document(self, file_path: str) -> List[Document]:
        """Load a document using basic file reading"""
        ext = Path(file_path).suffix.lower()
        
        if ext not in config.ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
        
        logger.info(f"Loading document: {file_path} (type: {ext})")
        
        try:
            if ext == '.pdf':
                return self._load_pdf(file_path)
            elif ext == '.csv':
                return self._load_csv(file_path)
            elif ext in ['.docx', '.doc']:
                return self._load_docx(file_path)
            else:  # .txt, .md, etc.
                return self._load_text(file_path)
        except Exception as e:
            logger.error(f"Failed to load {file_path}: {e}")
            # Fallback to text loading
            return self._load_text(file_path)
    
    def _load_text(self, file_path: str) -> List[Document]:
        """Load text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                metadata = {
                    "source": Path(file_path).name,
                    "file_path": file_path
                }
                return [Document(page_content=content, metadata=metadata)]
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
                metadata = {
                    "source": Path(file_path).name,
                    "file_path": file_path
                }
                return [Document(page_content=content, metadata=metadata)]
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load PDF file using PyPDF2"""
        try:
            import PyPDF2
            documents = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    text = page.extract_text()
                    if text.strip():
                        metadata = {
                            "source": Path(file_path).name,
                            "file_path": file_path,
                            "page": page_num
                        }
                        documents.append(Document(page_content=text, metadata=metadata))
            return documents
        except ImportError:
            logger.warning("PyPDF2 not available, using text extraction fallback")
            return self._load_text(file_path)
    
    def _load_csv(self, file_path: str) -> List[Document]:
        """Load CSV file - optimized for large files"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8', newline='') as file:
                # Try to detect dialect
                sample = file.read(1024)
                file.seek(0)
                
                try:
                    dialect = csv.Sniffer().sniff(sample)
                    has_header = csv.Sniffer().has_header(sample)
                except:
                    dialect = csv.excel
                    has_header = True
                
                if has_header:
                    reader = csv.DictReader(file, dialect=dialect)
                    for i, row in enumerate(reader, start=1):
                        # Limit row processing for very large files
                        if i > 1000:  # Process max 1000 rows to avoid huge vectors
                            logger.warning(f"CSV file too large, processing first 1000 rows only")
                            break
                            
                        content = "\n".join([f"{key}: {value}" for key, value in row.items() if value])
                        metadata = {
                            "source": Path(file_path).name,
                            "row": i,
                            "file_path": file_path
                        }
                        documents.append(Document(page_content=content, metadata=metadata))
                else:
                    reader = csv.reader(file, dialect=dialect)
                    for i, row in enumerate(reader, start=1):
                        if i > 1000:
                            logger.warning(f"CSV file too large, processing first 1000 rows only")
                            break
                            
                        content = " | ".join([f"Column {idx}: {value}" for idx, value in enumerate(row, start=1) if value])
                        metadata = {
                            "source": Path(file_path).name,
                            "row": i,
                            "file_path": file_path
                        }
                        documents.append(Document(page_content=content, metadata=metadata))
        except Exception as e:
            logger.error(f"CSV loading failed for {file_path}: {e}")
            # Fallback to text
            return self._load_text(file_path)
        
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Load DOCX file using python-docx"""
        try:
            import docx
            documents = []
            doc = docx.Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            metadata = {
                "source": Path(file_path).name,
                "file_path": file_path
            }
            return [Document(page_content=content, metadata=metadata)]
        except ImportError:
            logger.warning("python-docx not available, using text extraction fallback")
            return self._load_text(file_path)

class VectorStoreManager:
    """Manages vector store operations"""
    
    def __init__(self):
        logger.info("Initializing VectorStoreManager...")
        
        # Initialize embeddings
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-mpnet-base-v2",
                model_kwargs={'device': 'cpu'}
            )
            logger.info("Embeddings initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize embeddings: {e}")
            raise
        
        # Initialize vector store
        try:
            self.vector_store = Chroma(
                collection_name="rag_collection",
                embedding_function=self.embeddings,
                persist_directory=config.CHROMA_PERSIST_DIR,
            )
            logger.info("Vector store initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize vector store: {e}")
            raise
        
        self.document_loader = DocumentLoader()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # Reduced for better performance
            chunk_overlap=50,
            length_function=len,
        )
    
    def add_files(self, file_paths: List[str]) -> List[dict]:
        """Add files to vector store with detailed results"""
        results = []
        all_chunks = []
        
        for file_path in file_paths:
            result = {
                'filename': Path(file_path).name,
                'status': 'success',
                'chunks_created': 0,
                'error': None
            }
            
            try:
                # Check if file exists
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"File not found: {file_path}")
                
                logger.info(f"Starting to process: {file_path}")
                
                # Load document
                documents = self.document_loader.load_document(file_path)
                
                if not documents:
                    result['status'] = 'failed'
                    result['error'] = 'No documents loaded'
                    results.append(result)
                    continue
                
                logger.info(f"Loaded {len(documents)} documents from {file_path}")
                
                # Split into chunks
                chunks = self.text_splitter.split_documents(documents)
                all_chunks.extend(chunks)
                
                result['chunks_created'] = len(chunks)
                logger.info(f"Processed {file_path}: {len(documents)} documents -> {len(chunks)} chunks")
                
            except Exception as e:
                result['status'] = 'failed'
                result['error'] = str(e)
                logger.error(f"Failed to process {file_path}: {e}")
            
            results.append(result)
        
        # Add all chunks to vector store in batch
        if all_chunks:
            try:
                logger.info(f"Starting to add {len(all_chunks)} chunks to vector store...")
                
                # Add in smaller batches to avoid timeouts
                batch_size = 100
                for i in range(0, len(all_chunks), batch_size):
                    batch = all_chunks[i:i + batch_size]
                    self.vector_store.add_documents(batch)
                    logger.info(f"Added batch {i//batch_size + 1}/{(len(all_chunks)-1)//batch_size + 1}")
                    time.sleep(0.1)  # Small delay to prevent overwhelming
                
                logger.info(f"Successfully added {len(all_chunks)} chunks to vector store")
                
                # Verify the documents were added
                try:
                    doc_count = self.vector_store._collection.count()
                    logger.info(f"Vector store now contains {doc_count} documents")
                except Exception as e:
                    logger.warning(f"Could not verify document count: {e}")
                    
            except Exception as e:
                logger.error(f"Failed to add documents to vector store: {e}")
                for result in results:
                    if result['status'] == 'success':
                        result['status'] = 'failed'
                        result['error'] = f"Failed to add to vector store: {str(e)}"
        
        return results
    
    def reset_and_add_files(self, file_paths: List[str]) -> List[dict]:
        """Clear vector store and add new files"""
        logger.info("Resetting vector store...")
        
        try:
            # Clear existing collection
            self.vector_store.delete_collection()
            logger.info("Cleared existing vector store collection")
        except Exception as e:
            logger.warning(f"Failed to clear collection: {e}")
        
        # Reinitialize vector store
        try:
            self.vector_store = Chroma(
                collection_name="rag_collection",
                embedding_function=self.embeddings,
                persist_directory=config.CHROMA_PERSIST_DIR,
            )
            logger.info("Reinitialized vector store")
        except Exception as e:
            logger.error(f"Failed to reinitialize vector store: {e}")
            raise
        
        return self.add_files(file_paths)
    
    def similarity_search(self, query: str, k: int = None) -> List[Document]:
        """Search for similar documents"""
        if k is None:
            k = config.SIMILARITY_TOP_K
        
        try:
            logger.info(f"Performing similarity search for: '{query}' (k={k})")
            
            # First, check if we have any documents
            try:
                doc_count = self.vector_store._collection.count()
                logger.info(f"Vector store has {doc_count} documents")
                if doc_count == 0:
                    logger.warning("Vector store is empty - no documents to search")
                    return []
            except Exception as e:
                logger.warning(f"Could not check document count: {e}")
            
            results = self.vector_store.similarity_search(query, k=k)
            logger.info(f"Found {len(results)} relevant documents")
            
            # Log what we found for debugging
            for i, doc in enumerate(results):
                logger.info(f"Result {i+1}: {doc.metadata.get('source', 'Unknown')} - {doc.page_content[:100]}...")
                
            return results
        except Exception as e:
            logger.error(f"Similarity search failed: {e}")
            return []
    
    def get_collection_info(self) -> dict:
        """Get information about the vector store collection"""
        try:
            # Try to get document count
            doc_count = self.vector_store._collection.count()
            return {
                'document_count': doc_count,
                'persist_directory': config.CHROMA_PERSIST_DIR,
                'status': 'healthy'
            }
        except Exception as e:
            logger.error(f"Failed to get collection info: {e}")
            return {
                'document_count': 0,
                'persist_directory': config.CHROMA_PERSIST_DIR,
                'status': 'error',
                'error': str(e)
            }

# Global instance
vector_store_manager = VectorStoreManager()